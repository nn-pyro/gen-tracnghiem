import streamlit as st
import pdfplumber
import docx
import os
from langchain_together import ChatTogether
from langchain_core.prompts import ChatPromptTemplate

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text


def extract_text_from_word(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_text_from_txt(file):
    text = file.read().decode("utf-8")
    return text


def create_word(questions, title):
    doc = docx.Document()
    
    
    doc.add_heading(title.replace("##", "").strip(), level=1)

    for question in questions:
        question = question.strip()
        # Loại bỏ dấu ** và ## khỏi câu hỏi
        question = question.replace("**", "").replace("##", "").strip()
        if question:
            doc.add_paragraph(question)


    word_file_path = f"{title}.docx"
    doc.save(word_file_path)

    return word_file_path


llm = ChatTogether(
    model="google/gemma-2-9b-it",
    api_key="1d8fc48b1d2113e4a511c6309c08b9881a3e4a537476c829838089229b9f544f",
    temperature = 0.5,
    max_new_tokens=5120,
)


def display_title():
    st.title("Tạo trắc nghiệm")
    st.markdown("""
    <style>
        .title {
            text-align: center;
            font-size: 30px;
            color: #4A4A4A;
            margin-bottom: 20px;
        }
    </style>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    display_title()

    uploaded_file = st.file_uploader("Chọn file (PDF, DOCX, TXT):", type=["pdf", "docx", "txt"])

    if uploaded_file is not None:
        if uploaded_file.name.endswith('.pdf'):
            text_data = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.endswith('.docx'):
            text_data = extract_text_from_word(uploaded_file)
        elif uploaded_file.name.endswith('.txt'):
            text_data = extract_text_from_txt(uploaded_file)
        else:
            st.error("Định dạng file không được hỗ trợ. Vui lòng tải lên file PDF, DOCX hoặc TXT.")
            text_data = ""

        if text_data:
            num_questions = st.number_input("Số câu hỏi muốn tạo:", min_value=1, max_value=100, value=40)

            
            col1, col2 = st.columns(2)

            with col1:
                if st.button("Tạo Bộ Đề Trắc Nghiệm Có Kèm Đáp Án"):
                    with st.spinner("Đang tạo bộ đề trắc nghiệm..."):
                        quiz_prompt = ChatPromptTemplate.from_messages(
                            [
                                ("system", "Bạn là một trợ lý hữu ích giúp tạo câu hỏi trắc nghiệm từ văn bản. Bạn chỉ được phép tạo bộ đề trắc nghiệm với 4 đáp án. Không được phép diễn giải thêm điều gì. Hãy tạo đúng số lượng câu hỏi mà người dùng yêu cầu."),
                                ("human", f"Dựa trên văn bản sau, hãy tạo một bộ {num_questions} câu hỏi trắc nghiệm với bốn lựa chọn (A, B, C, D) và chỉ rõ đáp án đúng:\n\n{text_data}")
                            ]
                        )
                        
                        messages = quiz_prompt.format_messages()

                        quiz_response = llm(messages)

                        
                        st.write("Bộ Đề Trắc Nghiệm Có Kèm Đáp Án:")
                        questions = []
                        for question in quiz_response.content.split("\n"):
                            if question.strip(): 
                                questions.append(question)

                        
                        word_file_path = create_word(questions, "Bộ Đề Trắc Nghiệm Có Kèm Đáp Án")
                        with open(word_file_path, "rb") as word_file:
                            st.download_button("Tải xuống", data=word_file, file_name="bodetracnghiem.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        os.remove("./Bộ Đề Trắc Nghiệm Có Kèm Đáp Án.docx")

            with col2:
                if st.button("Tạo Bộ Đề Trắc Nghiệm Không Kèm Đáp Án"):
                    with st.spinner("Đang tạo bộ đề trắc nghiệm..."):
                        quiz_prompt_no_answer = ChatPromptTemplate.from_messages(
                            [
                                ("system", "Bạn là một trợ lý hữu ích giúp tạo câu hỏi trắc nghiệm từ văn bản. Bạn chỉ được phép tạo bộ đề trắc nghiệm với 4 đáp án. Không được phép diễn giải thêm điều gì. Hãy tạo đúng số lượng câu hỏi mà người dùng yêu cầu."),
                                ("human", f"Dựa trên văn bản sau, hãy tạo một bộ {num_questions} câu hỏi trắc nghiệm với bốn lựa chọn (A, B, C, D) mà không cần chỉ rõ đáp án đúng:\n\n{text_data}")
                            ]
                        )
                        
                        messages_no_answer = quiz_prompt_no_answer.format_messages()

                        quiz_response_no_answer = llm(messages_no_answer)

                        
                        st.write("Bộ Đề Trắc Nghiệm Không Kèm Đáp Án:")
                        questions_no_answer = []
                        for question in quiz_response_no_answer.content.split("\n"):
                            if question.strip():
                                questions_no_answer.append(question)

                        
                        word_file_path_no_answer = create_word(questions_no_answer, "Bộ Đề Trắc Nghiệm Không Kèm Đáp Án")
                        with open(word_file_path_no_answer, "rb") as word_file_no_answer:
                            st.download_button("Tải xuống", data=word_file_no_answer, file_name="bodetracnghiem.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        os.remove("./Bộ Đề Trắc Nghiệm Không Kèm Đáp Án.docx")
            
            st.markdown("---")
            st.subheader("Nội dung đã tạo:")
            if 'questions' in locals():
                for q in questions:
                    st.write(q)

            if 'questions_no_answer' in locals():
                for q in questions_no_answer:
                    st.write(q)
